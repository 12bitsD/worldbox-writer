from __future__ import annotations

import json
import re
from datetime import datetime, timezone
from html import escape
from io import BytesIO
from typing import Any, Dict, Iterable, List, Sequence, Tuple

from worldbox_writer.core.models import WorldState

DOCX_MIME_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)
PDF_MIME_TYPE = "application/pdf"
EXPORT_ARTIFACT_KINDS = {
    "novel_txt",
    "novel_markdown",
    "novel_html",
    "novel_docx",
    "novel_pdf",
    "world_settings_json",
    "timeline_json",
    "manifest_json",
}


def build_export_bundle(
    sim_id: str,
    branch_id: str,
    world: WorldState,
    nodes: Sequence[Dict[str, Any]],
) -> Dict[str, Any]:
    generated_at = datetime.now(timezone.utc).isoformat()
    bundle_name = _bundle_name(world.title, sim_id, branch_id)
    story_sections = _build_story_sections(nodes)
    world_settings = _build_world_settings(world)
    timeline = _build_timeline(nodes)
    novel = _build_plain_text(world, story_sections)
    markdown = _build_markdown(world, story_sections)
    html = _build_html(
        sim_id=sim_id,
        world=world,
        branch_id=branch_id,
        story_sections=story_sections,
        world_settings=world_settings,
        timeline=timeline,
        generated_at=generated_at,
    )

    return {
        "sim_id": sim_id,
        "branch_id": branch_id,
        "generated_at": generated_at,
        "summary": {
            "node_count": len(nodes),
            "rendered_node_count": len(story_sections),
            "character_count": len(world.characters),
            "rule_count": len(world.world_rules),
            "faction_count": len(world.factions),
            "location_count": len(world.locations),
        },
        "manifest": {
            "bundle_name": bundle_name,
            "generated_at": generated_at,
            "sim_id": sim_id,
            "branch_id": branch_id,
            "files": [
                _manifest_item(
                    "novel_txt", f"{bundle_name}.txt", "text/plain;charset=utf-8"
                ),
                _manifest_item(
                    "novel_markdown",
                    f"{bundle_name}.md",
                    "text/markdown;charset=utf-8",
                ),
                _manifest_item(
                    "novel_html",
                    f"{bundle_name}.html",
                    "text/html;charset=utf-8",
                ),
                _manifest_item("novel_docx", f"{bundle_name}.docx", DOCX_MIME_TYPE),
                _manifest_item("novel_pdf", f"{bundle_name}.pdf", PDF_MIME_TYPE),
                _manifest_item(
                    "world_settings_json",
                    f"{bundle_name}-settings.json",
                    "application/json",
                ),
                _manifest_item(
                    "timeline_json",
                    f"{bundle_name}-timeline.json",
                    "application/json",
                ),
                _manifest_item(
                    "manifest_json",
                    f"{bundle_name}-manifest.json",
                    "application/json",
                ),
            ],
        },
        "story_sections": story_sections,
        "novel": novel,
        "markdown": markdown,
        "html": html,
        "world_settings": world_settings,
        "timeline": timeline,
    }


def render_export_artifact(bundle: Dict[str, Any], kind: str) -> Tuple[str, str, bytes]:
    if kind not in EXPORT_ARTIFACT_KINDS:
        raise ValueError(f"Unsupported export artifact kind: {kind}")

    filename = _manifest_filename(bundle, kind)

    if kind == "novel_txt":
        return filename, "text/plain;charset=utf-8", bundle["novel"].encode("utf-8")
    if kind == "novel_markdown":
        return (
            filename,
            "text/markdown;charset=utf-8",
            bundle["markdown"].encode("utf-8"),
        )
    if kind == "novel_html":
        return filename, "text/html;charset=utf-8", bundle["html"].encode("utf-8")
    if kind == "novel_docx":
        return filename, DOCX_MIME_TYPE, _build_docx_bytes(bundle)
    if kind == "novel_pdf":
        return filename, PDF_MIME_TYPE, _build_pdf_bytes(bundle)
    if kind == "world_settings_json":
        return filename, "application/json", _json_bytes(bundle["world_settings"])
    if kind == "timeline_json":
        return filename, "application/json", _json_bytes(bundle["timeline"])
    return filename, "application/json", _json_bytes(bundle["manifest"])


def _manifest_item(kind: str, filename: str, mime_type: str) -> Dict[str, str]:
    return {"kind": kind, "filename": filename, "mime_type": mime_type}


def _manifest_filename(bundle: Dict[str, Any], kind: str) -> str:
    for item in bundle["manifest"]["files"]:
        if item["kind"] == kind:
            return str(item["filename"])
    raise ValueError(f"Missing manifest entry for export artifact kind: {kind}")


def _bundle_name(title: str, sim_id: str, branch_id: str) -> str:
    slug = re.sub(r"[^\w\u4e00-\u9fff-]+", "-", title.strip()).strip("-")
    return f"{slug or 'worldbox-story'}-{sim_id}-{branch_id}"


def _build_story_sections(nodes: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    sections: List[Dict[str, Any]] = []
    for node in nodes:
        prose = str(node.get("rendered_text") or "").strip()
        editor_html = str(node.get("editor_html") or "").strip() or None
        if not prose and not editor_html:
            continue
        sections.append(
            {
                "tick": int(node.get("tick", 0)),
                "title": str(node.get("title") or "未命名节点"),
                "type": str(node.get("node_type") or "development"),
                "description": str(node.get("description") or ""),
                "branch_id": str(node.get("branch_id") or "main"),
                "prose": prose,
                "editor_html": editor_html,
            }
        )
    return sections


def _build_world_settings(world: WorldState) -> Dict[str, Any]:
    return {
        "title": world.title,
        "premise": world.premise,
        "world_rules": list(world.world_rules),
        "factions": list(world.factions),
        "locations": list(world.locations),
        "characters": [
            {
                "name": character.name,
                "description": character.description,
                "personality": character.personality,
                "goals": list(character.goals),
                "status": character.status.value,
            }
            for character in world.characters.values()
        ],
    }


def _build_timeline(nodes: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    return [
        {
            "tick": int(node.get("tick", 0)),
            "title": str(node.get("title") or "未命名节点"),
            "type": str(node.get("node_type") or "development"),
            "description": str(node.get("description") or ""),
            "intervention": node.get("intervention_instruction"),
            "branch_id": str(node.get("branch_id") or "main"),
        }
        for node in nodes
    ]


def _build_plain_text(
    world: WorldState, story_sections: Sequence[Dict[str, Any]]
) -> str:
    sections = [f"{world.title}\n{'=' * 40}", f"前提：{world.premise}"]
    for section in story_sections:
        sections.append(f"【{section['title']}】\n\n{section['prose']}")
    if len(sections) == 2:
        return "\n\n".join(sections)
    body = ("\n\n" + "-" * 40 + "\n\n").join(sections[2:])
    return f"{sections[0]}\n\n{sections[1]}\n\n{'=' * 40}\n\n{body}"


def _build_markdown(world: WorldState, story_sections: Sequence[Dict[str, Any]]) -> str:
    lines = [f"# {world.title}", "", f"> {world.premise}", ""]
    for section in story_sections:
        lines.extend([f"## {section['title']}", "", section["prose"], ""])
    return "\n".join(lines).strip() + "\n"


def _build_html(
    *,
    sim_id: str,
    world: WorldState,
    branch_id: str,
    story_sections: Sequence[Dict[str, Any]],
    world_settings: Dict[str, Any],
    timeline: Sequence[Dict[str, Any]],
    generated_at: str,
) -> str:
    story_markup = (
        "\n".join(_render_story_section(section) for section in story_sections)
        or "<p class='muted'>暂无可导出的正文。</p>"
    )
    rules = (
        "".join(
            f"<li>{escape(str(rule))}</li>" for rule in world_settings["world_rules"]
        )
        or "<li>暂无世界规则</li>"
    )
    factions = _render_entity_cards(world_settings["factions"], "暂无势力设定。")
    locations = _render_entity_cards(world_settings["locations"], "暂无地点设定。")
    characters = (
        "".join(
            (
                "<article class='card'>"
                f"<h3>{escape(str(character['name']))}</h3>"
                f"<p>{escape(str(character.get('description') or character['personality']))}</p>"
                f"<p class='muted'>目标：{escape(' / '.join(character['goals']))}</p>"
                f"<p class='muted'>状态：{escape(str(character['status']))}</p>"
                "</article>"
            )
            for character in world_settings["characters"]
        )
        or "<p class='muted'>暂无角色档案。</p>"
    )
    timeline_rows = (
        "".join(
            (
                "<tr>"
                f"<td>T{item['tick']}</td>"
                f"<td>{escape(str(item['title']))}</td>"
                f"<td>{escape(str(item['type']))}</td>"
                f"<td>{escape(str(item['description']))}</td>"
                "</tr>"
            )
            for item in timeline
        )
        or "<tr><td colspan='4'>暂无时间线</td></tr>"
    )

    return f"""<!doctype html>
<html lang="zh-CN">
  <head>
    <meta charset="utf-8" />
    <meta name="viewport" content="width=device-width, initial-scale=1" />
    <title>{escape(world.title)}</title>
    <style>
      :root {{
        color-scheme: light;
        --ink: #1f1d1a;
        --muted: #6a635b;
        --paper: #f7f2e8;
        --panel: #fffdf8;
        --border: #d6c8b3;
      }}
      * {{ box-sizing: border-box; }}
      body {{
        margin: 0;
        font-family: "Georgia", "Songti SC", serif;
        color: var(--ink);
        background:
          radial-gradient(circle at top, rgba(123, 77, 42, 0.12), transparent 42%),
          linear-gradient(180deg, #f8f4eb 0%, var(--paper) 100%);
      }}
      main {{
        max-width: 920px;
        margin: 0 auto;
        padding: 48px 24px 72px;
      }}
      header {{
        padding-bottom: 28px;
        margin-bottom: 28px;
        border-bottom: 1px solid var(--border);
      }}
      h1, h2, h3 {{ margin: 0 0 12px; }}
      h1 {{ font-size: 44px; line-height: 1.1; }}
      h2 {{ font-size: 24px; margin-top: 32px; }}
      p, li, td {{ font-size: 16px; line-height: 1.9; }}
      .lede {{ font-size: 18px; color: var(--muted); max-width: 72ch; }}
      .meta {{ font-size: 13px; color: var(--muted); }}
      .story-section {{
        padding: 20px 0;
        border-bottom: 1px dashed rgba(123, 77, 42, 0.2);
      }}
      .card-grid {{
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(220px, 1fr));
        gap: 16px;
      }}
      .card {{
        background: var(--panel);
        border: 1px solid var(--border);
        padding: 16px;
        border-radius: 14px;
      }}
      .muted {{ color: var(--muted); }}
      table {{
        width: 100%;
        border-collapse: collapse;
        background: var(--panel);
        border: 1px solid var(--border);
      }}
      th, td {{
        padding: 10px 12px;
        border-bottom: 1px solid rgba(214, 200, 179, 0.8);
        text-align: left;
        vertical-align: top;
      }}
      @media print {{
        body {{ background: #fff; }}
        main {{ max-width: none; padding: 0; }}
      }}
    </style>
  </head>
  <body>
    <main>
      <header>
        <h1>{escape(world.title)}</h1>
        <p class="lede">{escape(world.premise)}</p>
        <p class="meta">Simulation {escape(sim_id)} | Branch {escape(branch_id)} | Generated at {escape(generated_at)}</p>
      </header>

      <section>
        <h2>正文</h2>
        {story_markup}
      </section>

      <section>
        <h2>世界规则</h2>
        <ul>{rules}</ul>
      </section>

      <section>
        <h2>角色设定</h2>
        <div class="card-grid">{characters}</div>
      </section>

      <section>
        <h2>势力设定</h2>
        <div class="card-grid">{factions}</div>
      </section>

      <section>
        <h2>地点设定</h2>
        <div class="card-grid">{locations}</div>
      </section>

      <section>
        <h2>故事时间线</h2>
        <table>
          <thead>
            <tr>
              <th>Tick</th>
              <th>标题</th>
              <th>类型</th>
              <th>描述</th>
            </tr>
          </thead>
          <tbody>{timeline_rows}</tbody>
        </table>
      </section>
    </main>
  </body>
</html>
"""


def _build_docx_bytes(bundle: Dict[str, Any]) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt

    document = Document()
    document.core_properties.title = bundle["world_settings"]["title"]
    document.core_properties.subject = bundle["world_settings"]["premise"]
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    title = document.add_heading(bundle["world_settings"]["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    preface = document.add_paragraph(bundle["world_settings"]["premise"])
    preface.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        f"Branch: {bundle['branch_id']} | Generated at: {bundle['generated_at']}"
    )

    for section in bundle["story_sections"]:
        document.add_heading(section["title"], level=1)
        document.add_paragraph(
            f"Tick {section['tick']} | {section['type']} | {section['branch_id']}"
        )
        for paragraph in _paragraphs(section["prose"]):
            document.add_paragraph(paragraph)

    document.add_page_break()
    document.add_heading("世界设定", level=1)
    document.add_heading("世界规则", level=2)
    for rule in bundle["world_settings"]["world_rules"] or ["暂无世界规则"]:
        document.add_paragraph(str(rule), style="List Bullet")

    document.add_heading("角色设定", level=2)
    for character in bundle["world_settings"]["characters"]:
        document.add_paragraph(character["name"], style="List Bullet")
        document.add_paragraph(
            str(character.get("description") or character["personality"])
        )
        document.add_paragraph(f"目标：{' / '.join(character['goals'])}")
        document.add_paragraph(f"状态：{character['status']}")

    document.add_heading("势力设定", level=2)
    for entity in bundle["world_settings"]["factions"] or [
        {"name": "暂无势力设定", "description": ""}
    ]:
        document.add_paragraph(str(entity.get("name") or "未命名"), style="List Bullet")
        if entity.get("description"):
            document.add_paragraph(str(entity["description"]))

    document.add_heading("地点设定", level=2)
    for entity in bundle["world_settings"]["locations"] or [
        {"name": "暂无地点设定", "description": ""}
    ]:
        document.add_paragraph(str(entity.get("name") or "未命名"), style="List Bullet")
        if entity.get("description"):
            document.add_paragraph(str(entity["description"]))

    document.add_heading("故事时间线", level=1)
    for item in bundle["timeline"]:
        document.add_paragraph(
            f"T{item['tick']} · {item['title']} [{item['type']}]",
            style="List Number",
        )
        document.add_paragraph(item["description"])

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _build_pdf_bytes(bundle: Dict[str, Any]) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import (
        PageBreak,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    font_name = "STSong-Light"
    if font_name not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ExportTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=24,
        leading=30,
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    meta_style = ParagraphStyle(
        "ExportMeta",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=10,
        leading=14,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#666666"),
        spaceAfter=8,
    )
    heading_style = ParagraphStyle(
        "ExportHeading",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=16,
        leading=22,
        spaceBefore=14,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "ExportBody",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=11,
        leading=18,
        spaceAfter=8,
    )

    flowables: List[Any] = [
        Paragraph(escape(bundle["world_settings"]["title"]), title_style),
        Paragraph(escape(bundle["world_settings"]["premise"]), meta_style),
        Paragraph(
            escape(
                f"Branch: {bundle['branch_id']} | Generated at: {bundle['generated_at']}"
            ),
            meta_style,
        ),
        Spacer(1, 8 * mm),
    ]

    for section in bundle["story_sections"]:
        flowables.append(Paragraph(escape(section["title"]), heading_style))
        flowables.append(
            Paragraph(
                escape(
                    f"Tick {section['tick']} | {section['type']} | {section['branch_id']}"
                ),
                meta_style,
            )
        )
        for paragraph in _paragraphs(section["prose"]):
            flowables.append(
                Paragraph(escape(paragraph).replace("\n", "<br/>"), body_style)
            )

    flowables.append(PageBreak())
    flowables.extend(
        [
            Paragraph("世界规则", heading_style),
            *[
                Paragraph(f"• {escape(str(rule))}", body_style)
                for rule in bundle["world_settings"]["world_rules"] or ["暂无世界规则"]
            ],
            Paragraph("角色设定", heading_style),
        ]
    )
    for character in bundle["world_settings"]["characters"]:
        flowables.append(Paragraph(f"• {escape(character['name'])}", body_style))
        flowables.append(
            Paragraph(
                escape(str(character.get("description") or character["personality"])),
                body_style,
            )
        )
        flowables.append(
            Paragraph(escape(f"目标：{' / '.join(character['goals'])}"), body_style)
        )
        flowables.append(Paragraph(escape(f"状态：{character['status']}"), body_style))

    flowables.append(Paragraph("故事时间线", heading_style))
    timeline_rows = [["Tick", "标题", "类型", "描述"]]
    for item in bundle["timeline"]:
        timeline_rows.append(
            [
                f"T{item['tick']}",
                str(item["title"]),
                str(item["type"]),
                str(item["description"]),
            ]
        )
    table = Table(timeline_rows, colWidths=[18 * mm, 42 * mm, 28 * mm, 88 * mm])
    table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#efe4d2")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#bca88e")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    flowables.append(table)

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=bundle["world_settings"]["title"],
        author="WorldBox Writer",
    )
    document.build(flowables)
    return buffer.getvalue()


def _json_bytes(payload: Any) -> bytes:
    return json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")


def _paragraphs(text: str) -> List[str]:
    parts = [segment.strip() for segment in text.splitlines() if segment.strip()]
    return parts or ["（正文为空）"]


def _render_entity_cards(entities: Sequence[Dict[str, Any]], empty_text: str) -> str:
    cards = "".join(
        (
            "<article class='card'>"
            f"<h3>{escape(str(entity.get('name') or '未命名'))}</h3>"
            f"<p>{escape(str(entity.get('description') or ''))}</p>"
            "</article>"
        )
        for entity in entities
    )
    return cards or f"<p class='muted'>{escape(empty_text)}</p>"


def _render_story_section(section: Dict[str, Any]) -> str:
    title = escape(str(section["title"]))
    editor_html = str(section.get("editor_html") or "").strip()
    if editor_html:
        body = editor_html
    else:
        body = "\n".join(
            f"<p>{escape(paragraph)}</p>" for paragraph in _paragraphs(section["prose"])
        )
    return f"<article class='story-section'><h3>{title}</h3>{body}</article>"
